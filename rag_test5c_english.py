import streamlit as st
import os
import faiss
import numpy as np
import sqlite3
import glob
import re
import requests

from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
from mistralai import Mistral
from dotenv import load_dotenv
from ragas.run_config import RunConfig
from datetime import datetime


load_dotenv()


os.environ["OPENAI_API_KEY"] = "sk-dummy-key-for-ragas"
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# Now import OpenAI dependencies
from datasets import Dataset
from ragas import evaluate
from langchain_mistralai import ChatMistralAI
from langchain_deepseek import ChatDeepSeek
import pandas as pd
import random

from ragas.metrics.collections import faithfulness, answer_relevancy, context_precision, context_recall
from langchain_huggingface import HuggingFaceEmbeddings



# ======================
# ⚙️ SETTINGS - MUST BE FIRST!
# ======================
st.set_page_config(page_title="RAG Assistant", layout="wide")

load_dotenv()

os.environ["TOKENIZERS_PARALLELISM"] = "false"

MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

if not MISTRAL_API_KEY:
    st.error("❌ MISTRAL_API_KEY not found in .env file!")
    st.stop()

EMB_MODEL = "all-MiniLM-L6-v2"

# Default values
DEFAULT_CHUNK_SIZE = 500
DEFAULT_CHUNK_OVERLAP = 100
DEFAULT_TOP_K = 30
DEFAULT_TEMPERATURE = 0.0
DEFAULT_LLM_MODEL = "mistral-large-latest"
DEFAULT_LLM_PROVIDER = "Mistral"
DEFAULT_PROMPT = "Strict RAG (Factual)"

# ======================
# PROMPT TEMPLATES (4 prompts)
# ======================

PROMPT_TEMPLATES = {
    "Strict RAG (Factual)": {
        "system_prompt": """You are a strict assistant. 
The material you see is news or texts from various sources.

CRITICAL:
- Answer ONLY if the information exists in the context
- DO NOT make assumptions
- If there is no answer, say:
  "There is not enough information"

Context:
{context}

Sources:
{sources}

Question:
{query}

Answer:""",
        "description": "Strict, factual answers only from the data"
    },
    
    "Journalistic Style (Generative)": {
        "system_prompt": """You are an experienced journalist working for the same media outlet from which the texts originate. 
Your style is professional, fluid, and journalistic.

INSTRUCTIONS:
- Use the context to write a short report or analysis
- Maintain the tone and style of the specific media outlet
- You can make connections between different pieces of information
- Add titles and subtitles where appropriate
- Do NOT add information that is NOT in the context

Context:
{context}

Sources:
{sources}

Topic/Question:
{query}

Report:""",
        "description": "Create reports and texts with journalistic style"
    },
    
    "Analysis & Key Points": {
        "system_prompt": """You are a news analyst. 
The material you see is news or texts from various sources.

INSTRUCTIONS:
- Extract the 3-5 most important points from the context
- Present them in bullet points
- Add a brief summary (1-2 sentences) at the beginning
- If there is numerical data, highlight it
- Do NOT add information that is NOT in the context

Context:
{context}

Sources:
{sources}

Topic/Question:
{query}

Analysis:""",
        "description": "Extract key points and analysis"
    },
    
    "Archivist (Documentation & Citations)": {
        "system_prompt": """You are an archivist and documentarian. 
Your job is to identify, cite, and accurately document information from archives/texts.

INSTRUCTIONS:
1. For each piece of information you cite, REFERENCE the source (filename/title)
2. QUOTE the exact passage from the text
3. IDENTIFY which file/location the information is in
4. If the concept/topic exists in multiple files, PRESENT all of them
5. Provide a CONCISE description of the content before the citations
6. ORGANIZE the answer into categories (e.g., "In file X...", "Additionally in...")
7. If the information is not found, say: "The specific concept/topic was not found in the archives"

FORMATTING:
- **File/Source:** [filename]
- **Quote:** "[exact text]"
- **Location:** [page/chunk if available]

Context:
{context}

Sources:
{sources}

Topic/Concept to search:
{query}

ARCHIVE / DOCUMENTATION:""",
        "description": "Identification, citations, and exact quotes from files"
    }
}

# ======================
# 🤖 MODELS
# ======================

@st.cache_resource
def load_models():
    model = SentenceTransformer(EMB_MODEL)
    
    # Mistral client
    mistral_client = Mistral(api_key=MISTRAL_API_KEY)
    
    # DeepSeek client (OpenAI-compatible)
    deepseek_client = None
    if DEEPSEEK_API_KEY:
        from openai import OpenAI
        deepseek_client = OpenAI(
            api_key=DEEPSEEK_API_KEY,
            base_url="https://api.deepseek.com/v1"
        )
    
    return model, mistral_client, deepseek_client

model, mistral_client, deepseek_client = load_models()

# ======================
# 📄 TEXT EXTRACTION
# ======================

def extract_text_from_pdf(pdf_file):
    try:
        if isinstance(pdf_file, str):
            # If it's a string (path), open the file
            with open(pdf_file, "rb") as f:
                reader = PdfReader(f)
        else:
            # If it's bytes (uploaded file), create a file-like object
            import io
            reader = PdfReader(io.BytesIO(pdf_file))  # ← FIXED
        
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"PDF error: {e}")
        return ""


def extract_text_from_txt(txt_source):
    try:
        if isinstance(txt_source, str):
            with open(txt_source, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        else:
            return txt_source.decode("utf-8", errors="ignore").strip()
    except Exception as e:
        st.error(f"TXT error: {e}")
        return ""

def extract_text_from_docx(docx_source):
    try:
        from docx import Document
        if isinstance(docx_source, str):
            doc = Document(docx_source)
        else:
            import io
            doc = Document(io.BytesIO(docx_source))
        return "\n".join([para.text for para in doc.paragraphs])
    except ImportError:
        st.warning("⚠️ For DOCX you need `pip install python-docx`")
        return ""
    except Exception as e:
        st.error(f"DOCX error: {e}")
        return ""

def extract_text_from_bytes(file_bytes, filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".txt" or ext == ".md":
        return extract_text_from_txt(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    else:
        return None

def extract_text_from_file_path(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".txt" or ext == ".md":
        return extract_text_from_txt(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        return None

# ======================
# ✂️ CHUNKING
# ======================

def split_text(text, chunk_size, overlap):
    chunks = []
    start = 0
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start += chunk_size - overlap
    
    chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
    return chunks

# ======================
# 🧠 FAISS + DOCS
# ======================

def build_index_from_texts(texts_with_meta, model):
    if not texts_with_meta:
        return None, None
    
    texts = [t["text"] for t in texts_with_meta]
    embeddings = model.encode(texts, convert_to_numpy=True, normalize_embeddings=True)
    
    if embeddings.ndim == 1:
        embeddings = embeddings.reshape(1, -1)
    
    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings.astype(np.float32))
    
    return index, texts_with_meta

# ======================
# 🗄️ DB
# ======================

def load_db_data(db_path, chunk_size, chunk_overlap):
    if not os.path.exists(db_path):
        return []
    
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("SELECT title, date, url, text, summary FROM news_table")
    rows = cursor.fetchall()
    conn.close()

    docs = []
    for row in rows:
        title = row[0]
        date = row[1]
        url = row[2]
        text = row[3]
        summary = row[4]
        if not text:
            continue
        full_text = f"{title}\n\n{text}\n\nSummary: {summary}"
        chunks = split_text(full_text, chunk_size, chunk_overlap)
        for chunk in chunks:
            docs.append({
                "text": chunk,
                "metadata": {
                    "title": title,
                    "date": date,
                    "url": url
                }
            })
    return docs

def load_or_build_db(chunk_size, chunk_overlap):
    docs = load_db_data("DBs/news_05.db", chunk_size, chunk_overlap)
    if not docs:
        return None, None
    index, docs_out = build_index_from_texts(docs, model)
    return index, docs_out

# ======================
# 📁 FOLDER
# ======================

def load_folder_data(folder_path, chunk_size, chunk_overlap):
    docs = []
    supported = [".pdf", ".txt", ".md", ".docx"]
    
    for ext in supported:
        for file_path in glob.glob(os.path.join(folder_path, f"*{ext}"), recursive=False):
            filename = os.path.basename(file_path)
            text = extract_text_from_file_path(file_path)
            if not text:
                continue
            chunks = split_text(text, chunk_size, chunk_overlap)
            for i, chunk in enumerate(chunks):
                docs.append({
                    "text": chunk,
                    "metadata": {
                        "title": filename,
                        "date": "local_file",
                        "url": file_path,
                        "chunk_id": i
                    }
                })
    return docs

# ======================
# 🔍 RAG QUERY (with prompt template)
# ======================

def query_rag_with_prompt(index, docs, query, model, client, top_k, temperature, llm_model, provider, prompt_key):
    """Executes RAG query with selected prompt template"""
    
    # 1. Retrieve chunks
    q_emb = model.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
    actual_top_k = min(top_k, index.ntotal)
    D, I = index.search(q_emb, actual_top_k)
    retrieved = [docs[i] for i in I[0] if i < len(docs)]
    
    # 2. Create context and sources
    context = "\n\n".join([d["text"] for d in retrieved])
    sources = "\n".join([
        f"{d['metadata']['title']} ({d.get('metadata',{}).get('date','N/A')})"
        for d in retrieved
    ])
    
    # 3. Select prompt template
    template = PROMPT_TEMPLATES[prompt_key]
    prompt = template["system_prompt"].format(
        context=context,
        sources=sources,
        query=query
    )
    
    # 4. Call LLM
    if provider == "Mistral":
        response = client.chat.complete(
            model=llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature
        )
        return response.choices[0].message.content
    else:  # DeepSeek
        response = client.chat.completions.create(
            model=llm_model,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature
        )
        return response.choices[0].message.content

# ======================
# 🎯 RAGAS EVALUATION (new component)
# ======================

def create_ragas_dataset(test_questions, ground_truths, index, docs, model, client, provider, llm_model, top_k, temperature, prompt_key):
    """
    Creates dataset for RAGAS evaluation using the existing RAG system
    """
    results = {
        "question": [],
        "answer": [],
        "contexts": [],
        "ground_truth": []
    }
    
    for i, (query, gt) in enumerate(zip(test_questions, ground_truths)):
        # 1. Retrieve chunks (same code as query_rag_with_prompt)
        q_emb = model.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
        actual_top_k = min(top_k, index.ntotal)
        D, I = index.search(q_emb, actual_top_k)
        retrieved = [docs[i] for i in I[0] if i < len(docs)]
        
        # 2. Create context (the retrieved chunks)
        contexts = [d["text"] for d in retrieved]
        
        # 3. Create answer (calls your LLM)
        context_str = "\n\n".join(contexts)
        sources = "\n".join([d['metadata']['title'] for d in retrieved])
        
        template = PROMPT_TEMPLATES[prompt_key]
        prompt = template["system_prompt"].format(
            context=context_str,
            sources=sources,
            query=query
        )
        
        if provider == "Mistral":
            response = client.chat.complete(
                model=llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=temperature
            )
            answer = response.choices[0].message.content
        else:
            response = client.chat.completions.create(
                model=llm_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=temperature
            )
            answer = response.choices[0].message.content
        
        # 4. Store
        results["question"].append(query)
        results["answer"].append(answer)
        results["contexts"].append(contexts)  # RAGAS expects list of strings
        results["ground_truth"].append(gt)
        
        # Progress update
        st.write(f"✅ Question {i+1}/{len(test_questions)} completed")
    
    return results


def evaluate_with_ragas(test_questions, ground_truths, index, docs, model, client, provider, llm_model, top_k, temperature, prompt_key):
    """
    Executes RAGAS evaluation using the existing LLM as judge
    """
    try:
        from datasets import Dataset
        from ragas import evaluate
        from ragas.metrics import faithfulness, answer_relevancy, context_precision, context_recall
        from langchain_mistralai import ChatMistralAI
        from langchain_deepseek import ChatDeepSeek
    except ImportError:
        st.error("❌ Need to install: pip install ragas datasets langchain-mistralai langchain-deepseek")
        return None
    
    # 1. Create dataset from RAG system
    with st.spinner("🔄 Running RAG system for each question..."):
        results = create_ragas_dataset(
            test_questions, ground_truths, index, docs, model, 
            client, provider, llm_model, top_k, temperature, prompt_key
        )
    
    dataset = Dataset.from_dict(results)
    
    # 2. Use THE SAME LLM as judge (Mistral or DeepSeek)
    with st.spinner("⚖️ Evaluating with RAGAS..."):
        if provider == "Mistral":
            evaluator_llm = ChatMistralAI(
                model=llm_model,
                mistral_api_key=MISTRAL_API_KEY,
                temperature=0
            )
        else:
            evaluator_llm = ChatDeepSeek(
                model=llm_model,
                api_key=DEEPSEEK_API_KEY,
                temperature=0
            )
        
        score = evaluate(
            dataset,
            metrics=[faithfulness, answer_relevancy, context_precision, context_recall],
            llm=evaluator_llm
        )
    
    return score, dataset



# ======================
# 🖥️ UI WITH TABS (Updated)
# ======================

# Initialize session_state
if "index" not in st.session_state:
    st.session_state.index = None
if "docs" not in st.session_state:
    st.session_state.docs = None
if "source_name" not in st.session_state:
    st.session_state.source_name = None
if "current_chunk_size" not in st.session_state:
    st.session_state.current_chunk_size = DEFAULT_CHUNK_SIZE
if "current_chunk_overlap" not in st.session_state:
    st.session_state.current_chunk_overlap = DEFAULT_CHUNK_OVERLAP
if "current_top_k" not in st.session_state:
    st.session_state.current_top_k = DEFAULT_TOP_K
if "current_temperature" not in st.session_state:
    st.session_state.current_temperature = DEFAULT_TEMPERATURE
if "current_llm_model" not in st.session_state:
    st.session_state.current_llm_model = DEFAULT_LLM_MODEL
if "current_llm_provider" not in st.session_state:
    st.session_state.current_llm_provider = DEFAULT_LLM_PROVIDER
if "current_prompt" not in st.session_state:
    st.session_state.current_prompt = DEFAULT_PROMPT

# Initialize for evaluation
if "test_questions" not in st.session_state:
    st.session_state.test_questions = []
if "test_ground_truths" not in st.session_state:
    st.session_state.test_ground_truths = []
if "test_sources" not in st.session_state:
    st.session_state.test_sources = []
if "selected_chunk_idx" not in st.session_state:
    st.session_state.selected_chunk_idx = None
if "generated_qa" not in st.session_state:
    st.session_state.generated_qa = None
if "active_tab" not in st.session_state:
    st.session_state.active_tab = "RAG Assistant"

# Create Tabs (4 tabs now)
tab1, tab2, tab3, tab4 = st.tabs(["📚 RAG Assistant", "📊 Evaluation", "📝 Create Test Set", "ℹ️ About"])

# ======================
# TAB 1: RAG ASSISTANT (Existing)
# ======================
with tab1:
    st.title("📚 RAG Assistant")
    
    # Two columns
    left_col, right_col = st.columns([4, 1], gap="medium")
    
    # ========== LEFT COLUMN ==========
    with left_col:
        st.header("🔍 Knowledge Source Selection")
        
        source_option = st.radio(
            "Choose source:",
            ["🗄️ SQLite Database (news_05.db)", "📂 Local Folder", "📄 Single File Upload"],
            horizontal=True,
            key="source_radio"
        )
        
        st.divider()
        
        # ========== CHUNKING PARAMETERS ==========
        st.subheader("⚙️ Chunking Settings")
        
        col_size, col_overlap = st.columns(2)
        with col_size:
            chunk_size = st.slider(
                "📏 Chunk Size (characters)",
                min_value=100,
                max_value=2000,
                value=st.session_state.current_chunk_size,
                step=50,
                key="chunk_size_slider"
            )
        with col_overlap:
            chunk_overlap = st.slider(
                "🔄 Chunk Overlap (characters)",
                min_value=0,
                max_value=500,
                value=st.session_state.current_chunk_overlap,
                step=25,
                key="chunk_overlap_slider"
            )
        
        # ========== LLM PARAMETERS ==========
        st.subheader("🎨 LLM Settings")
        
        # Provider selection
        provider_options = ["Mistral"]
        if deepseek_client is not None:
            provider_options.append("DeepSeek")
        
        selected_provider = st.selectbox(
            "🤖 Select Provider",
            provider_options,
            index=0 if st.session_state.current_llm_provider == "Mistral" else 1,
            key="provider_select"
        )
        
        # Model selection based on provider
        if selected_provider == "Mistral":
            model_options = {
                "mistral-large-latest": "Mistral Large (latest)",
                "mistral-small-latest": "Mistral Small (latest)",
                "open-mistral-nemo": "Mistral NeMo",
            }
        else:  # DeepSeek
            model_options = {
                "deepseek-chat": "DeepSeek-V3.2 (Chat)",
                "deepseek-reasoner": "DeepSeek-R1 (Reasoner)",
            }
        
        selected_model = st.selectbox(
            "🧠 Select model",
            options=list(model_options.keys()),
            format_func=lambda x: model_options[x],
            key="model_select"
        )
        
        temperature = st.slider(
            "🌡️ Temperature",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.current_temperature,
            step=0.05,
            format="%.2f",
            key="temperature_slider"
        )
        
        # ========== PROMPT TEMPLATE SELECTION ==========
        st.subheader("📝 Response Style")
        
        prompt_option = st.selectbox(
            "Choose response style:",
            options=list(PROMPT_TEMPLATES.keys()),
            format_func=lambda x: f"{x} - {PROMPT_TEMPLATES[x]['description']}",
            key="prompt_select"
        )
        
        # Update session state
        if selected_provider != st.session_state.current_llm_provider:
            st.session_state.current_llm_provider = selected_provider
        
        if selected_model != st.session_state.current_llm_model:
            st.session_state.current_llm_model = selected_model
        
        if temperature != st.session_state.current_temperature:
            st.session_state.current_temperature = temperature
        
        if prompt_option != st.session_state.current_prompt:
            st.session_state.current_prompt = prompt_option
        
        # Check if chunking parameters changed
        if chunk_size != st.session_state.current_chunk_size or chunk_overlap != st.session_state.current_chunk_overlap:
            st.session_state.current_chunk_size = chunk_size
            st.session_state.current_chunk_overlap = chunk_overlap
            st.session_state.index = None
            st.session_state.docs = None
            st.session_state.source_name = None
            st.info("⚙️ Chunking parameters changed - you need to reload the data")
        
        st.divider()
        
        # 1️⃣ SQLite DB
        if source_option == "🗄️ SQLite Database (news_05.db)":
            st.subheader("🗄️ SQLite Database")
            st.info("Loading database from DBs/news_05.db")
            
            if st.button("📥 Load DB", type="primary", key="load_db"):
                with st.spinner(f"Loading from DB with chunk_size={chunk_size}, overlap={chunk_overlap}..."):
                    index, docs = load_or_build_db(chunk_size, chunk_overlap)
                    if index is not None:
                        st.session_state.index = index
                        st.session_state.docs = docs
                        st.session_state.source_name = f"SQLite Database (size={chunk_size}, overlap={chunk_overlap})"
                        st.success(f"✅ Loaded {len(docs)} chunks from DB")
                        st.rerun()
                    else:
                        st.error("❌ DB not found or empty")
        
        # 2️⃣ LOCAL FOLDER
        elif source_option == "📂 Local Folder":
            st.subheader("📂 Local Folder")
            
            folder_path = st.text_input("Folder path:", placeholder="e.g., /Users/username/Documents/my_files", key="folder_path")
            
            if st.button("📁 Load folder", type="primary", key="load_folder"):
                if not folder_path or not os.path.isdir(folder_path):
                    st.error("❌ Invalid path!")
                else:
                    with st.spinner(f"Reading files from {folder_path}..."):
                        docs = load_folder_data(folder_path, chunk_size, chunk_overlap)
                        if docs:
                            index, _ = build_index_from_texts(docs, model)
                            st.session_state.index = index
                            st.session_state.docs = docs
                            st.session_state.source_name = f"Folder: {folder_path} (size={chunk_size}, overlap={chunk_overlap})"
                            st.success(f"✅ Loaded {len(docs)} chunks from {folder_path}")
                            st.rerun()
                        else:
                            st.error("❌ No supported files found")
        
        # 3️⃣ SINGLE FILE UPLOAD
        else:
            st.subheader("📄 Single File Upload")
            
            uploaded_file = st.file_uploader(
                "Choose file:",
                type=["pdf", "txt", "md", "docx"],
                help="Supported: PDF, TXT, MD, DOCX",
                key="file_uploader"
            )
            
            if st.button("📄 Load file", type="primary", key="load_file"):
                if uploaded_file is not None:
                    with st.spinner(f"Processing {uploaded_file.name}..."):
                        text = extract_text_from_bytes(uploaded_file.getvalue(), uploaded_file.name)
                        if text:
                            chunks = split_text(text, chunk_size, chunk_overlap)
                            docs = []
                            for i, chunk in enumerate(chunks):
                                docs.append({
                                    "text": chunk,
                                    "metadata": {
                                        "title": uploaded_file.name,
                                        "date": "uploaded_file",
                                        "filename": uploaded_file.name,
                                        "chunk_id": i
                                    }
                                })
                            index, _ = build_index_from_texts(docs, model)
                            st.session_state.index = index
                            st.session_state.docs = docs
                            st.session_state.source_name = f"File: {uploaded_file.name} (size={chunk_size}, overlap={chunk_overlap})"
                            st.success(f"✅ Loaded {len(docs)} chunks from {uploaded_file.name}")
                            st.rerun()
                        else:
                            st.error("❌ Could not extract text")
                else:
                    st.warning("⚠️ Please choose a file first.")
        
        st.divider()
        
        # 💬 QUESTION
        if st.session_state.index is not None and st.session_state.docs is not None:
            st.header("💬 Ask your question")
            
            # Top K slider
            top_k = st.slider(
                "🔍 Number of chunks to retrieve (Top K)",
                min_value=1,
                max_value=min(500, len(st.session_state.docs)),
                value=min(st.session_state.current_top_k, len(st.session_state.docs)),
                step=10,
                key="top_k_slider"
            )
            st.session_state.current_top_k = top_k
            
            query = st.text_area("Question:", height=100, placeholder="Type your question here...", key="query_input")
            
            ask_button = st.button("🚀 Submit question", type="primary", use_container_width=True, key="ask_btn")
            
            if ask_button and query.strip():
                provider = st.session_state.current_llm_provider
                llm_model = st.session_state.current_llm_model
                prompt_key = st.session_state.current_prompt
                
                # Select client
                if provider == "Mistral":
                    client = mistral_client
                else:
                    client = deepseek_client
                
                with st.spinner(f"Searching with {provider}, {prompt_key}, top_k={top_k}, temperature={temperature}..."):
                    try:
                        answer = query_rag_with_prompt(
                            st.session_state.index, 
                            st.session_state.docs, 
                            query, 
                            model, 
                            client,
                            top_k,
                            temperature,
                            llm_model,
                            provider,
                            prompt_key
                        )
                        st.markdown("### 🧠 Answer:")
                        st.write(answer)
                    except Exception as e:
                        st.error(f"Error: {e}")
            elif ask_button and not query.strip():
                st.warning("⚠️ Please write a question first.")
        else:
            st.info("👈 Choose a source above, adjust parameters, and load data first.")
    
    # ========== RIGHT COLUMN (Permanent) ==========
    with right_col:
        st.header("ℹ️ Information")
        
        st.info("📝 **Instructions:**\n\n1. Adjust chunk size/overlap\n2. Select LLM provider & model\n3. Choose response style\n4. Adjust temperature\n5. Load data\n6. Ask questions!")
        
        st.divider()
        
        st.subheader("📊 System Status")
        
        if st.session_state.source_name:
            st.metric("📌 Active source", st.session_state.source_name)
        else:
            st.metric("📌 Active source", "None", delta="⚠️ Not loaded")
        
        if st.session_state.docs:
            st.metric("📄 Number of chunks", len(st.session_state.docs))
        else:
            st.metric("📄 Number of chunks", "0", delta="📭 Load data")
        
        st.divider()
        st.subheader("🔧 Active Parameters")
        st.metric("🤖 LLM Provider", st.session_state.current_llm_provider)
        st.metric("🧠 Model", st.session_state.current_llm_model.split("-")[0] if "-" in st.session_state.current_llm_model else st.session_state.current_llm_model[:20])
        st.metric("📏 Chunk Size", f"{st.session_state.current_chunk_size} chars")
        st.metric("🔄 Overlap", f"{st.session_state.current_chunk_overlap} chars")
        st.metric("🔍 Top K", st.session_state.current_top_k)
        st.metric("🌡️ Temperature", f"{st.session_state.current_temperature:.2f}")
        st.metric("📝 Style", st.session_state.current_prompt.split(" ")[0])
        
        st.divider()
        
        with st.expander("📝 Prompt Descriptions"):
            for key, value in PROMPT_TEMPLATES.items():
                st.markdown(f"**{key}**")
                st.caption(value["description"])
                st.markdown("---")
        
        with st.expander("🔧 Technical Information"):
            st.write(f"**Embeddings model:** {EMB_MODEL}")
            st.write(f"**FAISS index type:** IndexFlatIP")
            if st.session_state.current_llm_provider == "Mistral":
                st.write("**LLM:** Mistral AI (API)")
            else:
                st.write("**LLM:** DeepSeek (API)")
        
        with st.expander("💡 Tips"):
            st.markdown("""
            **Response Styles:**
            - **Strict RAG:** For precise answers
            - **Journalistic:** For reports/articles
            - **Analysis:** For key points
            - **Archivist:** For citations & quotes
            
            **Archivist examples:**
            - "Which files contain the word 'X'"
            - "Find all quotes about the economy"
            - "Cite the prime minister's statements"
            
            **Chunk Size:**
            - Small (200-300): Precision
            - Large (800-1000): Context
            """)


# ======================
# TAB 2: EVALUATION (WITH CHUNK SIZE & OVERLAP + GT SIMILARITY)
# ======================
with tab2:
    st.header("📊 RAG System Evaluation")
    st.markdown("Evaluation using embeddings - Save and compare results")
    
    if st.session_state.index is None or st.session_state.docs is None:
        st.warning("⚠️ You need to load data in the RAG Assistant tab first")
    elif not st.session_state.test_questions:
        st.warning("⚠️ No test set available. Go to 'Create Test Set' tab")
    else:
        # ========== CHUNKING PARAMETERS (NEW!) ==========
        st.subheader("✂️ Chunking Parameters")
        st.caption("These parameters affect HOW texts are split into chunks")
        
        col_chunk1, col_chunk2 = st.columns(2)
        with col_chunk1:
            eval_chunk_size = st.slider(
                "📏 Chunk Size (characters)",
                min_value=100,
                max_value=2000,
                value=st.session_state.current_chunk_size,
                step=50,
                help="Size of each chunk in characters. Smaller = more precision, larger = more context",
                key="eval_chunk_size"
            )
        
        with col_chunk2:
            eval_chunk_overlap = st.slider(
                "🔄 Chunk Overlap (characters)",
                min_value=0,
                max_value=500,
                value=st.session_state.current_chunk_overlap,
                step=25,
                help="Overlap between consecutive chunks. Helps maintain continuous meaning",
                key="eval_chunk_overlap"
            )
        
        # Warning if parameters change
        if eval_chunk_size != st.session_state.current_chunk_size or eval_chunk_overlap != st.session_state.current_chunk_overlap:
            st.warning("⚠️ Chunking parameters changed. You need to reload the data in the RAG Assistant tab for them to apply.")
        
        st.divider()
        
        # ========== RETRIEVAL PARAMETERS ==========
        st.subheader("🔍 Retrieval Parameters")
        
        col_ret1, col_ret2 = st.columns(2)
        with col_ret1:
            eval_top_k = st.slider(
                "🔝 Top K chunks",
                min_value=1,
                max_value=min(50, len(st.session_state.docs)),
                value=min(10, len(st.session_state.docs)),
                help="Number of chunks to retrieve for each question",
                key="eval_top_k"
            )
        
        with col_ret2:
            similarity_threshold = st.slider(
                "🎯 Similarity Threshold",
                min_value=0.0,
                max_value=1.0,
                value=0.3,
                step=0.05,
                help="Minimum similarity to keep a chunk (0-1)"
            )
        
        st.divider()
        
        # ========== LLM PARAMETERS ==========
        st.subheader("🤖 LLM Parameters")
        
        col_llm1, col_llm2 = st.columns(2)
        
        with col_llm1:
            provider_options = ["Mistral"]
            if deepseek_client is not None:
                provider_options.append("DeepSeek")
            
            eval_provider = st.selectbox(
                "🤖 LLM Provider",
                provider_options,
                index=0 if st.session_state.current_llm_provider == "Mistral" else 1,
                key="eval_provider"
            )
        
        with col_llm2:
            if eval_provider == "Mistral":
                model_options = {
                    "mistral-large-latest": "Mistral Large",
                    "mistral-small-latest": "Mistral Small",
                    "open-mistral-nemo": "Mistral NeMo",
                }
            else:
                model_options = {
                    "deepseek-chat": "DeepSeek-V3.2",
                    "deepseek-reasoner": "DeepSeek-R1",
                }
            
            eval_model = st.selectbox(
                "🧠 Model",
                options=list(model_options.keys()),
                format_func=lambda x: model_options[x],
                key="eval_model"
            )
        
        col_temp1, col_temp2 = st.columns(2)
        with col_temp1:
            eval_temperature = st.slider(
                "🌡️ Temperature",
                min_value=0.0,
                max_value=1.0,
                value=st.session_state.current_temperature,
                step=0.05,
                format="%.2f",
                key="eval_temperature"
            )
        
        with col_temp2:
            eval_prompt = st.selectbox(
                "📝 Response Style",
                options=list(PROMPT_TEMPLATES.keys()),
                format_func=lambda x: f"{x} - {PROMPT_TEMPLATES[x]['description'][:50]}...",
                key="eval_prompt"
            )
        
        st.divider()
        
        # ========== RUN NAME ==========
        st.subheader("💾 Save Results")
        
        col_name1, col_name2 = st.columns([3, 1])
        with col_name1:
            run_name = st.text_input(
                "Run name",
                value=f"Run_{len(st.session_state.get('saved_runs', [])) + 1}",
                help="Give a name to identify this run"
            )
        
        with col_name2:
            save_results = st.checkbox("Save results", value=True)
        
        if "saved_runs" not in st.session_state:
            st.session_state.saved_runs = []
        
        st.divider()
        
        # ========== RUN EVALUATION ==========
        if st.button("🚀 Run Evaluation", type="primary", use_container_width=True):
            
            # Check if chunking parameters match loaded data
            if eval_chunk_size != st.session_state.current_chunk_size or eval_chunk_overlap != st.session_state.current_chunk_overlap:
                st.error("❌ Chunking parameters don't match loaded data. Go to RAG Assistant tab, change parameters and reload data.")
            else:
                with st.spinner("Evaluating..."):
                    try:
                        from sentence_transformers import util
                        import pandas as pd
                        import numpy as np
                        from datetime import datetime
                        
                        if eval_provider == "Mistral":
                            client = mistral_client
                        else:
                            client = deepseek_client
                        
                        results = []
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        for i, (query, gt) in enumerate(zip(st.session_state.test_questions, st.session_state.test_ground_truths)):
                            status_text.text(f"Processing {i+1}/{len(st.session_state.test_questions)}: {query[:50]}...")
                            
                            # 1. Retrieve chunks
                            q_emb = model.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
                            actual_top_k = min(eval_top_k, st.session_state.index.ntotal)
                            D, I = st.session_state.index.search(q_emb, actual_top_k)
                            
                            # Apply similarity threshold
                            retrieved = []
                            for idx, score in zip(I[0], D[0]):
                                if idx < len(st.session_state.docs) and score >= similarity_threshold:
                                    retrieved.append(st.session_state.docs[idx])
                            
                            if not retrieved:
                                retrieved = [st.session_state.docs[I[0][0]]] if len(I[0]) > 0 else []
                            
                            contexts = [d["text"] for d in retrieved]
                            context_str = "\n\n".join(contexts)
                            sources = "\n".join([d['metadata']['title'] for d in retrieved])
                            
                            # 2. Get answer from LLM
                            template = PROMPT_TEMPLATES[eval_prompt]
                            prompt = template["system_prompt"].format(
                                context=context_str, sources=sources, query=query
                            )
                            
                            if eval_provider == "Mistral":
                                response = client.chat.complete(
                                    model=eval_model,
                                    messages=[{"role": "user", "content": prompt}],
                                    temperature=eval_temperature
                                )
                                answer = response.choices[0].message.content
                            else:
                                response = client.chat.completions.create(
                                    model=eval_model,
                                    messages=[{"role": "user", "content": prompt}],
                                    temperature=eval_temperature
                                )
                                answer = response.choices[0].message.content
                            
                            # 3. Calculate metrics with embeddings
                            query_emb = model.encode(query, convert_to_numpy=True, normalize_embeddings=True)
                            answer_emb = model.encode(answer, convert_to_numpy=True, normalize_embeddings=True)
                            
                            # Context Precision
                            if contexts:
                                chunk_embs = model.encode(contexts, convert_to_numpy=True, normalize_embeddings=True)
                                similarities = util.cos_sim(query_emb, chunk_embs)[0]
                                context_precision = float(similarities.mean())
                            else:
                                context_precision = 0.0
                            
                            # Answer Relevance
                            answer_relevance = float(util.cos_sim(query_emb, answer_emb)[0][0])
                            
                            # Faithfulness
                            context_emb = model.encode(context_str[:3000], convert_to_numpy=True, normalize_embeddings=True)
                            faithfulness = float(util.cos_sim(answer_emb, context_emb)[0][0])
                            
                            # GT Similarity
                            gt_emb = model.encode(gt, convert_to_numpy=True, normalize_embeddings=True)
                            gt_similarity = float(util.cos_sim(answer_emb, gt_emb)[0][0]) if gt else 0
                            
                            results.append({
                                "question": query,
                                "answer": answer,
                                "ground_truth": gt,
                                "faithfulness": faithfulness,
                                "answer_relevance": answer_relevance,
                                "context_precision": context_precision,
                                "gt_similarity": gt_similarity,
                                "chunks_retrieved": len(contexts),
                                "chunks_total": actual_top_k
                            })
                            
                            progress_bar.progress((i + 1) / len(st.session_state.test_questions))
                        
                        status_text.text("✅ Completed!")
                        
                        # Create DataFrame
                        df_results = pd.DataFrame(results)
                        
                        # Calculate averages
                        avg_faithfulness = df_results["faithfulness"].mean()
                        avg_answer_relevance = df_results["answer_relevance"].mean()
                        avg_context_precision = df_results["context_precision"].mean()
                        avg_gt_similarity = df_results["gt_similarity"].mean()
                        
                        # ========== DISPLAY RESULTS ==========
                        st.success("✅ Evaluation completed!")
                        
                        # Metrics cards
                        st.subheader("📊 Metric Results")
                        
                        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                        with col_m1:
                            color = "🟢" if avg_faithfulness >= 0.7 else "🟡" if avg_faithfulness >= 0.4 else "🔴"
                            st.metric(f"{color} Faithfulness", f"{avg_faithfulness:.3f}")
                            st.caption("Answer faithfulness to context")
                        
                        with col_m2:
                            color = "🟢" if avg_answer_relevance >= 0.7 else "🟡" if avg_answer_relevance >= 0.4 else "🔴"
                            st.metric(f"{color} Answer Relevancy", f"{avg_answer_relevance:.3f}")
                            st.caption("Answer relevance to question")
                        
                        with col_m3:
                            color = "🟢" if avg_context_precision >= 0.7 else "🟡" if avg_context_precision >= 0.4 else "🔴"
                            st.metric(f"{color} Context Precision", f"{avg_context_precision:.3f}")
                            st.caption("Retrieved chunks quality")
                        
                        with col_m4:
                            color = "🟢" if avg_gt_similarity >= 0.7 else "🟡" if avg_gt_similarity >= 0.4 else "🔴"
                            st.metric(f"{color} GT Similarity", f"{avg_gt_similarity:.3f}")
                            st.caption("Similarity to correct answer")
                        
                        # ========== SAVE ==========
                        if save_results:
                            run_data = {
                                "name": run_name,
                                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                "parameters": {
                                    "chunk_size": eval_chunk_size,
                                    "chunk_overlap": eval_chunk_overlap,
                                    "top_k": eval_top_k,
                                    "similarity_threshold": similarity_threshold,
                                    "provider": eval_provider,
                                    "model": eval_model,
                                    "temperature": eval_temperature,
                                    "prompt": eval_prompt
                                },
                                "results": {
                                    "faithfulness": avg_faithfulness,
                                    "answer_relevance": avg_answer_relevance,
                                    "context_precision": avg_context_precision,
                                    "gt_similarity": avg_gt_similarity
                                },
                                "detailed_results": df_results.to_dict('records')
                            }
                            
                            st.session_state.saved_runs.append(run_data)
                            st.success(f"✅ Results saved as '{run_name}'")
                        
                        # ========== DETAILED RESULTS ==========
                        with st.expander("📊 Detailed results per question"):
                            # Display table with ALL metrics
                            display_df = df_results[["question", "faithfulness", "answer_relevance", "context_precision", "gt_similarity", "chunks_retrieved"]].copy()
                            display_df["question"] = display_df["question"].apply(lambda x: x[:80] + "..." if len(x) > 80 else x)
                            display_df.columns = ["Question", "Faithfulness", "Answer Relevancy", "Context Precision", "GT Similarity", "#Chunks"]
                            st.dataframe(display_df, use_container_width=True)
                            
                            # GT Similarity explanation
                            st.info("💡 **GT Similarity:** How similar the RAG answer is to the ground truth answer. Values > 0.7 indicate very good matching.")
                            
                            # Export
                            csv = df_results.to_csv(index=False).encode('utf-8')
                            st.download_button(
                                "📥 Download detailed results (CSV)",
                                csv,
                                f"evaluation_{run_name}.csv",
                                "text/csv"
                            )
                        
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        

        # ========== RESULTS HISTORY ==========
        if st.session_state.saved_runs:
            st.divider()
            st.subheader("📚 Results History")
            st.markdown("Compare different runs")
            
            comparison_data = []
            for run in st.session_state.saved_runs:
                comparison_data.append({
                    "Name": run["name"],
                    "Date": run["timestamp"],
                    # Metrics
                    "Faithfulness": f"{run['results']['faithfulness']:.3f}",
                    "Answer Rel.": f"{run['results']['answer_relevance']:.3f}",
                    "Context Prec.": f"{run['results']['context_precision']:.3f}",
                    "GT Similarity": f"{run['results']['gt_similarity']:.3f}",
                    # Chunking Parameters
                    "Chunk Size": run["parameters"].get("chunk_size", "N/A"),
                    "Chunk Overlap": run["parameters"].get("chunk_overlap", "N/A"),
                    # Retrieval Parameters
                    "Top K": run["parameters"].get("top_k", "N/A"),
                    "Threshold": run["parameters"].get("threshold", run["parameters"].get("similarity_threshold", "N/A")),
                    # LLM Parameters
                    "Provider": run["parameters"].get("provider", "N/A"),
                    "Model": run["parameters"].get("model", "N/A").split("-")[0] if "model" in run["parameters"] else "N/A",
                    "Temperature": run["parameters"].get("temperature", "N/A"),
                    "Prompt": run["parameters"].get("prompt", "N/A")[:30] + "..." if len(run["parameters"].get("prompt", "")) > 30 else run["parameters"].get("prompt", "N/A")
                })
            
            df_comparison = pd.DataFrame(comparison_data)
            
            # Display table with all columns
            st.dataframe(df_comparison, use_container_width=True)
            
            # Column selection for display (optional)
            with st.expander("🔧 Select columns to display"):
                all_columns = df_comparison.columns.tolist()
                selected_columns = st.multiselect(
                    "Choose which columns to display:",
                    all_columns,
                    default=all_columns[:8]  # Default first 8 columns
                )
                if selected_columns:
                    st.dataframe(df_comparison[selected_columns], use_container_width=True)
            
            # ========== CSV DOWNLOAD BUTTON ==========
            col_export1, col_export2 = st.columns(2)
            with col_export1:
                csv_history = df_comparison.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📥 Download history (CSV)",
                    data=csv_history,
                    file_name=f"evaluation_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col_export2:
                if st.button("🗑️ Clear all history", use_container_width=True):
                    st.session_state.saved_runs = []
                    st.rerun()
            
            # Select for details
            selected_run = st.selectbox(
                "Select run for details",
                options=[run["name"] for run in st.session_state.saved_runs]
            )
            
            if selected_run:
                run = next((r for r in st.session_state.saved_runs if r["name"] == selected_run), None)
                if run:
                    with st.expander(f"📋 Details for '{selected_run}'"):
                        # Display all parameters in a nice format
                        col1, col2 = st.columns(2)
                        with col1:
                            st.markdown("**🔧 Chunking & Retrieval Parameters:**")
                            st.json({
                                "Chunk Size": run["parameters"].get("chunk_size", "N/A"),
                                "Chunk Overlap": run["parameters"].get("chunk_overlap", "N/A"),
                                "Top K": run["parameters"].get("top_k", "N/A"),
                                "Threshold": run["parameters"].get("threshold", run["parameters"].get("similarity_threshold", "N/A"))
                            })
                        
                        with col2:
                            st.markdown("**🤖 LLM Parameters:**")
                            st.json({
                                "Provider": run["parameters"].get("provider", "N/A"),
                                "Model": run["parameters"].get("model", "N/A"),
                                "Temperature": run["parameters"].get("temperature", "N/A"),
                                "Prompt": run["parameters"].get("prompt", "N/A")
                            })
                        
                        # Metrics comparison chart
                        metrics_df = pd.DataFrame({
                            "Metric": ["Faithfulness", "Answer Relevancy", "Context Precision", "GT Similarity"],
                            "Value": [
                                run["results"]["faithfulness"],
                                run["results"]["answer_relevance"],
                                run["results"]["context_precision"],
                                run["results"]["gt_similarity"]
                            ]
                        })
                        st.bar_chart(metrics_df.set_index("Metric"))
                        
                        # Download detailed results
                        if "detailed_results" in run:
                            df_details = pd.DataFrame(run["detailed_results"])
                            st.download_button(
                                "📥 Download detailed results (CSV)",
                                df_details.to_csv(index=False).encode('utf-8'),
                                f"{selected_run}_details.csv",
                                "text/csv"
                            )
        
        # ========== METRICS EXPLANATION ==========
        st.divider()
        with st.expander("ℹ️ Metrics Explanation"):
            st.markdown("""
            **Faithfulness:** How much the answer is based on the retrieved context.  
            *Values close to 1.0 = Answer is faithful to the context*
            
            **Answer Relevancy:** How much the answer addresses the question.  
            *Values close to 1.0 = Answer is relevant*
            
            **Context Precision:** How relevant the retrieved chunks are.  
            *Values close to 1.0 = Chunks are relevant to the question*
            
            **GT Similarity:** How similar the RAG answer is to the ground truth answer.  
            *Values > 0.8 = Answer is almost identical*  
            *Values 0.5-0.8 = Similar meaning, different wording*  
            *Values < 0.5 = Answer differs significantly*
            
            **Chunking Parameters:**
            - **Chunk Size:** Larger = more context, smaller = more precision
            - **Chunk Overlap:** Helps maintain continuous meaning between chunks
            """)


# ======================
# TAB 3: CREATE TEST SET
# ======================
with tab3:
    st.header("📝 Create Test Set")
    st.markdown("Create questions and correct answers from your data")
    
    if st.session_state.index is None or st.session_state.docs is None:
        st.warning("⚠️ You need to load data in the RAG Assistant tab first")
    else:
        # 3 ways to create test set
        method = st.radio(
            "Choose creation method:",
            [
                "✍️ Manual (read and write)",
                "🤖 Automatic generation (with LLM)",
                "📁 Upload from CSV/JSON"
            ],
            horizontal=True
        )
        
        st.divider()
        
        # ========== METHOD 1: MANUAL ==========
        if method == "✍️ Manual (read and write)":
            st.subheader("📖 Select a text and create a question")
            
            # Select random or specific chunk
            select_mode = st.radio(
                "Text selection:",
                ["🔀 Random chunk", "📋 From chunk list", "🔍 Search by keyword"],
                horizontal=True
            )
            
            selected_text = None
            selected_metadata = None
            
            if select_mode == "🔀 Random chunk":
                if st.button("🎲 Select random text"):
                    import random
                    idx = random.randint(0, len(st.session_state.docs) - 1)
                    st.session_state.selected_chunk_idx = idx
                    st.rerun()
                
                if st.session_state.get("selected_chunk_idx") is not None:
                    idx = st.session_state.selected_chunk_idx
                    selected_text = st.session_state.docs[idx]["text"]
                    selected_metadata = st.session_state.docs[idx]["metadata"]
            
            elif select_mode == "📋 From chunk list":
                # Create list with preview
                chunk_options = []
                for i, doc in enumerate(st.session_state.docs[:100]):
                    preview = doc["text"][:80] + "..." if len(doc["text"]) > 80 else doc["text"]
                    chunk_options.append(f"{i}: {preview}")
                
                selected_idx = st.selectbox(
                    "Select chunk:",
                    options=range(len(chunk_options)),
                    format_func=lambda x: chunk_options[x]
                )
                selected_text = st.session_state.docs[selected_idx]["text"]
                selected_metadata = st.session_state.docs[selected_idx]["metadata"]
            
            else:  # Search
                search_term = st.text_input("🔍 Keyword:", placeholder="e.g., economy, growth, minister")
                if search_term:
                    results = []
                    for i, doc in enumerate(st.session_state.docs):
                        if search_term.lower() in doc["text"].lower():
                            preview = doc["text"][:100] + "..." if len(doc["text"]) > 100 else doc["text"]
                            results.append((i, preview))
                    
                    if results:
                        st.info(f"Found {len(results)} chunks")
                        selected_result = st.selectbox(
                            "Select result:",
                            options=results,
                            format_func=lambda x: f"{x[0]}: {x[1]}"
                        )
                        selected_text = st.session_state.docs[selected_result[0]]["text"]
                        selected_metadata = st.session_state.docs[selected_result[0]]["metadata"]
                    elif search_term:
                        st.warning("No results found")
            
            # Display selected text
            if selected_text:
                st.markdown("---")
                st.subheader("📄 Selected text:")
                
                # Metadata
                st.caption(f"**Source:** {selected_metadata.get('title', 'N/A')} | **Date:** {selected_metadata.get('date', 'N/A')}")
                
                # The text
                st.text_area("Text:", selected_text, height=200, key="preview_text")
                
                st.divider()
                
                # Question creation form
                st.subheader("✍️ Create question-answer")
                
                col_q, col_a = st.columns(2)
                with col_q:
                    new_question = st.text_area(
                        "📌 Question:",
                        placeholder="e.g., When was the article published?",
                        height=80,
                        key="new_question"
                    )
                with col_a:
                    new_answer = st.text_area(
                        "✅ Correct answer:",
                        placeholder="e.g., 2022",
                        height=80,
                        key="new_answer"
                    )
                
                # Optional note
                note = st.text_input("📝 Note (optional):", placeholder="e.g., Important detail...")
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("➕ Add to test set", use_container_width=True):
                        if new_question and new_answer:
                            st.session_state.test_questions.append(new_question)
                            st.session_state.test_ground_truths.append(new_answer)
                            st.session_state.test_sources.append({
                                "text": selected_text[:200],
                                "metadata": selected_metadata,
                                "note": note
                            })
                            st.success(f"✅ Added: {new_question[:50]}...")
                            st.rerun()
                        else:
                            st.warning("⚠️ Fill in both question and answer")
                
                with col_btn2:
                    if st.button("🗑️ Clear", use_container_width=True):
                        st.rerun()
        
        # ========== METHOD 2: AUTOMATIC GENERATION ==========
        elif method == "🤖 Automatic generation (with LLM)":
            st.subheader("🤖 Automatic question-answer generation")
            st.info("The LLM will read the texts and generate questions with answers")
            
            num_questions = st.slider("Number of questions to generate:", 1, 20, 5)
            
            # Select chunks to process
            use_all_chunks = st.checkbox("Use all chunks", value=False)
            
            if not use_all_chunks:
                num_chunks = st.slider("Number of chunks to process:", 1, 50, 10)
                chunks_to_process = st.session_state.docs[:num_chunks]
            else:
                chunks_to_process = st.session_state.docs
            
            if st.button("🚀 Generate questions", type="primary", use_container_width=True):
                with st.spinner(f"Generating {num_questions} questions with LLM..."):
                    # Prepare prompt
                    sample_chunks = chunks_to_process[:min(5, len(chunks_to_process))]
                    sample_texts = "\n\n---\n\n".join([c["text"][:500] for c in sample_chunks])
                    
                    generation_prompt = f"""Read the following texts and generate {num_questions} questions with their correct answers.

TEXTS:
{sample_texts}

BASIC RULES:
1. Each question must be answerable EXCLUSIVELY from the texts
2. The answer must be short and precise
3. Questions should be different from each other

FORMATTING (STRICT):
Question: [question text]
Answer: [answer text]
--- 
Question: [question text]
Answer: [answer text]
..."""

                    # Call LLM
                    provider = st.session_state.current_llm_provider
                    if provider == "Mistral":
                        response = mistral_client.chat.complete(
                            model=st.session_state.current_llm_model,
                            messages=[{"role": "user", "content": generation_prompt}],
                            temperature=0.3
                        )
                        generated = response.choices[0].message.content
                    else:
                        response = deepseek_client.chat.completions.create(
                            model=st.session_state.current_llm_model,
                            messages=[{"role": "user", "content": generation_prompt}],
                            temperature=0.3
                        )
                        generated = response.choices[0].message.content
                    
                    st.session_state.generated_qa = generated
                    st.success("✅ Generation completed!")
            
            # Display and save generated Q&A
            if st.session_state.get("generated_qa"):
                st.subheader("📋 Suggested questions-answers")
                
                # Simple parsing of output
                lines = st.session_state.generated_qa.split('\n')
                qa_pairs = []
                current_q = None
                current_a = None
                
                for line in lines:
                    if line.startswith("Question:"):
                        current_q = line.replace("Question:", "").strip()
                    elif line.startswith("Answer:"):
                        current_a = line.replace("Answer:", "").strip()
                        if current_q and current_a:
                            qa_pairs.append((current_q, current_a))
                            current_q = None
                            current_a = None
                
                # Display with checkboxes
                for i, (q, a) in enumerate(qa_pairs):
                    with st.container():
                        col_cb, col_q, col_a = st.columns([0.5, 3, 2])
                        with col_cb:
                            add = st.checkbox("➕", key=f"auto_add_{i}")
                        with col_q:
                            st.text_area("Question:", q, height=60, key=f"auto_q_{i}", disabled=not add)
                        with col_a:
                            st.text_area("Answer:", a, height=60, key=f"auto_a_{i}", disabled=not add)
                        
                        if add:
                            if st.button(f"Save {i+1}", key=f"save_auto_{i}"):
                                st.session_state.test_questions.append(q)
                                st.session_state.test_ground_truths.append(a)
                                st.session_state.test_sources.append({"auto_generated": True})
                                st.success(f"✅ Saved: {q[:50]}...")
                                st.rerun()
        
        # ========== METHOD 3: UPLOAD ==========
        else:  # Upload CSV/JSON
            st.subheader("📁 Upload existing test set")
            
            uploaded_test = st.file_uploader(
                "Upload file with questions and answers",
                type=["csv", "json", "xlsx"],
                help="File must have columns: question, ground_truth (and optionally: source, note)"
            )
            
            if uploaded_test:
                import pandas as pd
                
                try:
                    if uploaded_test.name.endswith('.csv'):
                        df = pd.read_csv(uploaded_test)
                    elif uploaded_test.name.endswith('.json'):
                        df = pd.read_json(uploaded_test)
                    else:  # xlsx
                        df = pd.read_excel(uploaded_test)
                    
                    st.success(f"✅ Loaded {len(df)} records")
                    st.dataframe(df.head())
                    
                    if st.button("📥 Replace current test set"):
                        st.session_state.test_questions = df['question'].tolist()
                        st.session_state.test_ground_truths = df['ground_truth'].tolist()
                        
                        if 'source' in df.columns:
                            st.session_state.test_sources = df['source'].tolist()
                        else:
                            st.session_state.test_sources = [{}] * len(df)
                        
                        st.success(f"✅ Replaced test set with {len(df)} questions")
                        st.rerun()
                except Exception as e:
                    st.error(f"Error reading file: {e}")
        
        # ========== DISPLAY CURRENT TEST SET ==========
        st.divider()
        st.subheader(f"📋 Current Test Set ({len(st.session_state.test_questions)} questions)")
        
        if st.session_state.test_questions:
            # Management options
            col_manage1, col_manage2, col_manage3 = st.columns(3)
            with col_manage1:
                if st.button("🗑️ Delete all", use_container_width=True):
                    st.session_state.test_questions = []
                    st.session_state.test_ground_truths = []
                    st.session_state.test_sources = []
                    st.rerun()
            
            with col_manage2:
                # Export to CSV
                import pandas as pd
                export_df = pd.DataFrame({
                    "question": st.session_state.test_questions,
                    "ground_truth": st.session_state.test_ground_truths
                })
                csv_data = export_df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📥 Export to CSV",
                    data=csv_data,
                    file_name="test_set.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col_manage3:
                st.info(f"✅ {len(st.session_state.test_questions)} questions ready for evaluation")
            
            # Display question list
            for i, (q, a) in enumerate(zip(st.session_state.test_questions, st.session_state.test_ground_truths)):
                with st.expander(f"{i+1}. {q[:80]}..."):
                    st.markdown(f"**✅ Correct answer:** {a}")
                    if i < len(st.session_state.test_sources):
                        src = st.session_state.test_sources[i]
                        if src and src.get("metadata"):
                            st.caption(f"**Source:** {src['metadata'].get('title', 'N/A')}")
                        if src and src.get("note"):
                            st.caption(f"**Note:** {src['note']}")
                    
                    # Delete button
                    if st.button(f"🗑️ Delete", key=f"delete_{i}"):
                        st.session_state.test_questions.pop(i)
                        st.session_state.test_ground_truths.pop(i)
                        if i < len(st.session_state.test_sources):
                            st.session_state.test_sources.pop(i)
                        st.rerun()
        else:
            st.info("👈 Create questions from your texts")


# ======================
# TAB 4: ABOUT
# ======================
with tab4:
    st.title("ℹ️ About RAG Assistant")
    st.markdown("""
    ### What is this application?
    
    It's a **RAG (Retrieval-Augmented Generation)** assistant that allows you to:
    
    - Load data from various sources (SQLite DB, local folder, or individual files)
    - **Test different chunking parameters in real time**
    - **Choose between Mistral and DeepSeek LLMs**
    - **Select between 4 different response styles**
    - **Evaluate the system with RAGAS metrics**
    - **Create test sets for objective comparison**
    
    ### The 4 response styles:
    
    1. **Strict RAG (Factual):** Answers only with existing information
    2. **Journalistic Style:** Writes reports like a journalist
    3. **Analysis & Key Points:** Extracts the most important points
    4. **Archivist:** Identifies, quotes, and documents with precision
    
    ### Evaluation Metrics:
    
    - **Faithfulness:** Is the answer based on the context?
    - **Answer Relevancy:** Is the answer relevant to the question?
    - **Context Precision:** Are the retrieved chunks relevant?
    - **Context Recall:** Were all necessary pieces of information found?
    
    ### Version: 7.0 (RAGAS Evaluation + Test Set Creator)
    """)

# ======================
# FOOTER
# ======================
st.divider()
st.caption("⚙️ RAG Assistant - 4 response styles + RAGAS Evaluation + Test Set Creator")